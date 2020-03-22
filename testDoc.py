from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import time
from datetime import datetime, date, timedelta

#start_date = str(date(2020, 3, 9))[5:].replace('-','/') 
start_date = date(2020, 3, 9)
end_date = start_date + timedelta(days=4)

document = Document()

def format_style(run, font = "Calibri", size = 18, style = None):
    run.font.name = font
    run.font.size = Pt(size)
    if(style != None):
        setattr(run, style, True)
    return

def date_to_string(date):
    return str(date)[5:].replace('-','/')


def create_Title(start_date, end_date):
    title = document.add_paragraph()
    title_text = title.add_run("ALG 1 - Lesson Plans - {0} - {1}".format(date_to_string(start_date), date_to_string(end_date)))
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 
    format_style(title_text, "Calibri", 18, 'underline')
    return

def create_day_schedule(date, item1, item2):
    heading = document.add_paragraph()
    style = heading.add_run(date)
    format_style(style, "Calibri", 18, 'underline')
    create_line_item(item1)
    create_line_item(item2)
    return

def create_line_item(item):
    line = document.add_paragraph()
    line.style = 'List Bullet'
    line_style = line.add_run(item)
    format_style(line_style)
    return

def weekday_num(date):
    day = date.weekday()
    if (day == 0):
        return "Monday "
    if (day == 1):
        return "Tuesday "
    if (day == 2):
        return "Wednesday "
    if (day == 3):
        return "Thursday "
    if (day == 4):
        return "Friday "

create_Title(start_date, end_date)
    
current_day = start_date
i = 0
while(i < 5):
    create_day_schedule(weekday_num(current_day)+date_to_string(current_day),
                        "This is Item 1", "This is Item 2")
    current_day += timedelta(days=1)
    i += 1

document.save("test.docx")
