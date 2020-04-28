from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import time
from datetime import datetime, date, timedelta

def format_style(run, font = "Calibri", size = 18, style = None):
    run.font.name = font
    run.font.size = Pt(size)
    if(style != None):
        setattr(run, style, True)
    return

def date_to_string(date):
    return str(date)[5:].replace('-','/')


def create_Title(document, section, start_date):
    title = document.add_paragraph()
    title_text = title.add_run("{0} Lesson Plans - {1} to {2}".format(section, start_date.strftime("%b %d"), (start_date + timedelta(days=4)).strftime("%b %d")))
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 
    format_style(title_text, "Calibri", 18, 'underline')
    return

def create_day_schedule(document, date, item1, item2):
    heading = document.add_paragraph()
    style = heading.add_run(date)
    format_style(style, "Calibri", 18, 'underline')
    create_line_item(document, item1)
    create_line_item(document, item2)
    return

def create_line_item(document, item):
    line = document.add_paragraph()
    line.style = 'List Bullet'
    line_style = line.add_run(item)
    format_style(line_style)
    return

def weekday_num(date):
    day = date.weekday()
    if (day == 0):
        return "Monday "
    if (day == 1):
        return "Tuesday "
    if (day == 2):
        return "Wednesday "
    if (day == 3):
        return "Thursday "
    if (day == 4):
        return "Friday "


def create_document(start_date, section, activity_list, homework_list):
    document = Document()
    start_date = start_date
    create_Title(document, section, start_date)
    current_day = start_date
    i = 0
    while(i < 5):
        create_day_schedule(document, current_day.strftime("%A %m/%d"),
                            activity_list.pop(0), homework_list.pop(0))
        current_day += timedelta(days=1)
        i += 1
    file_name = "Lesson Plans - {0} to {1}".format(start_date.strftime("%b %d"), (start_date + timedelta(days=4)).strftime("%b %d"))
    document.save(file_name + ".docx")
    return file_name